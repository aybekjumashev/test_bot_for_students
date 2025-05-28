# core/management/commands/create_test_data.py

import random
import os # Fayl nomlari bilan ishlash uchun
from django.core.management.base import BaseCommand
from django.utils.translation import gettext_lazy as _
from django.core.files.base import ContentFile
from io import BytesIO
from docx import Document as DocxDocument
from django.db import transaction

# Modellarni import qilish
from core.models import Subject, Question, EducationType # User hozircha shart emas

# Agar EducationType kerak bo'lsa, uni topish yoki yaratish
def get_or_create_default_education_type():
    obj, created = EducationType.objects.get_or_create(
        name_uz="Umumiy Ta'lim",
        defaults={
            'name_kaa': "Улыўма Би תлимлендириў", # Iltimos, to'g'ri qoraqalpoqcha matnini kiriting
            'name_ru': "Общее Образование",
            'is_otm': False # Yoki sizning ehtiyojingizga qarab
        }
    )
    if created:
        print(f"'{obj.name_uz}' education type created.")
    return obj

# Oddiy DOCX fayl kontentini generatsiya qilish uchun funksiya
def generate_dummy_docx_content(subject_name_display, question_num, language_display):
    document = DocxDocument()
    document.add_heading(f"{subject_name_display} ({language_display}) - Savol #{question_num}", level=1)
    document.add_paragraph(
        f"Bu {subject_name_display} fanidan {language_display} tilidagi {question_num}-savolning matni. "
        "Lorem ipsum dolor sit amet, consectetur adipiscing elit. "
        "Variantlar quyidagicha bo'lishi mumkin:"
    )
    document.add_paragraph("A) Variant A matni")
    document.add_paragraph("B) Variant B matni")
    document.add_paragraph("C) Variant C matni")
    document.add_paragraph("D) Variant D matni")

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

class Command(BaseCommand):
    help = 'Creates fake data for subjects and questions (multilingual) for testing purposes.'

    def add_arguments(self, parser):
        parser.add_argument(
            '--clear',
            action='store_true',
            help='Delete all existing Subjects and Questions before creating new ones.',
        )
        parser.add_argument(
            '--num_subjects',
            type=int,
            default=5,
            help='Number of subjects to create.',
        )
        parser.add_argument(
            '--questions_per_subject',
            type=int,
            default=10,
            help='Number of questions to create per subject.',
        )

    @transaction.atomic
    def handle(self, *args, **kwargs):
        clear_data = kwargs['clear']
        num_subjects_to_create = kwargs['num_subjects']
        questions_per_subject_to_create = kwargs['questions_per_subject']

        self.stdout.write(self.style.SUCCESS('Feyk ko\'p tilli ma\'lumotlarni yaratish boshlandi...'))

        if clear_data:
            Question.objects.all().delete()
            Subject.objects.all().delete()
            # EducationType.objects.all().delete() # Agar buni ham tozalash kerak bo'lsa
            self.stdout.write(self.style.WARNING('Mavjud Fanlar va Savollar o\'chirildi.'))

        # default_education_type = get_or_create_default_education_type() # Agar Subject modelida kerak bo'lsa

        # Fan nomlari uchun namunalar (ko'proq qo'shishingiz mumkin)
        subject_templates = [
            {'uz': "Matematika", 'kaa': "Математика", 'ru': "Математика"},
            {'uz': "Fizika", 'kaa': "Физика", 'ru': "Физика"},
            {'uz': "Ona Tili", 'kaa': "Ана Тили", 'ru': "Родной Язык"},
            {'uz': "Adabiyot", 'kaa': "Әдебият", 'ru': "Литература"},
            {'uz': "Tarix", 'kaa': "Тарийх", 'ru': "История"},
            {'uz': "Ingliz Tili", 'kaa': "Инглис Тили", 'ru': "Английский Язык"},
            {'uz': "Biologiya", 'kaa': "Биология", 'ru': "Биология"},
            {'uz': "Kimyo", 'kaa': "Химия", 'ru': "Химия"},
            {'uz': "Geografiya", 'kaa': "География", 'ru': "География"},
            {'uz': "Informatika", 'kaa': "Информатика", 'ru': "Информатика"},
        ]

        answer_choices = ['a', 'b', 'c', 'd']

        created_subjects_count = 0
        for i in range(num_subjects_to_create):
            s_template = subject_templates[i % len(subject_templates)] # Nomlarni aylantirib ishlatish
            
            subject_name_uz = s_template['uz'] + (f" {i // len(subject_templates) + 1}" if num_subjects_to_create > len(subject_templates) and i >= len(subject_templates) else "")
            subject_name_kaa = s_template['kaa'] + (f" {i // len(subject_templates) + 1}" if num_subjects_to_create > len(subject_templates) and i >= len(subject_templates) else "")
            subject_name_ru = s_template['ru'] + (f" {i // len(subject_templates) + 1}" if num_subjects_to_create > len(subject_templates) and i >= len(subject_templates) else "")

            subject, created = Subject.objects.get_or_create(
                name_uz=subject_name_uz, # Asosiy nom sifatida o'zbekchani olamiz (yoki boshqasini)
                defaults={
                    'name_kaa': subject_name_kaa,
                    'name_ru': subject_name_ru,
                    'min_course_year': 1,
                    'max_course_year': 11, # Yoki sizning ehtiyojingizga qarab
                    'is_active': True,
                    'voucher_template_name': subject_name_uz.replace(" ", "").replace("-","") # Masalan, "OnaTili"
                }
            )

            if created:
                self.stdout.write(self.style.SUCCESS(f'"{subject.name_uz}" fani yaratildi.'))
                created_subjects_count += 1
            else:
                self.stdout.write(self.style.NOTICE(f'"{subject.name_uz}" fani allaqachon mavjud.'))

            # Har bir fan uchun savollar yaratish
            current_question_count_for_subject = subject.questions.count()
            new_questions_needed = questions_per_subject_to_create - current_question_count_for_subject
            
            if new_questions_needed <= 0 and not created: # Agar fan avvaldan bor bo'lsa va savollari yetarli bo'lsa
                 self.stdout.write(self.style.NOTICE(f'"{subject.name_uz}" fani uchun savollar yetarli ({current_question_count_for_subject} ta).'))
                 continue
            elif new_questions_needed <=0 and created: # Agar fan yangi yaratilgan bo'lsa, lekin 0 ta savol so'ralgan bo'lsa
                pass # Hech nima qilmaymiz


            for j in range(new_questions_needed):
                question_instance_number = current_question_count_for_subject + j + 1
                
                # Har bir til uchun DOCX fayl yaratish
                docx_uz_stream = generate_dummy_docx_content(subject.name_uz, question_instance_number, "O'zbekcha")
                file_uz = ContentFile(docx_uz_stream.read(), name=f'q_uz_{subject.id}_{question_instance_number}.docx')
                
                docx_kaa_stream = generate_dummy_docx_content(subject.name_kaa, question_instance_number, "Qoraqalpoqcha")
                file_kaa = ContentFile(docx_kaa_stream.read(), name=f'q_kaa_{subject.id}_{question_instance_number}.docx')
                
                docx_ru_stream = generate_dummy_docx_content(subject.name_ru, question_instance_number, "Ruscha")
                file_ru = ContentFile(docx_ru_stream.read(), name=f'q_ru_{subject.id}_{question_instance_number}.docx')

                Question.objects.create(
                    subject=subject,
                    question_file_uz=file_uz,
                    question_file_kaa=file_kaa,
                    question_file_ru=file_ru,
                    correct_answer=random.choice(answer_choices),
                    is_active=True
                )
            if new_questions_needed > 0:
                self.stdout.write(self.style.SUCCESS(f'"{subject.name_uz}" fani uchun {new_questions_needed} ta yangi savol (3 tilda) yaratildi.'))

        if created_subjects_count == 0 and num_subjects_to_create > 0 and not clear_data:
             self.stdout.write(self.style.WARNING('Yangi fanlar yaratilmadi (ehtimol allaqachon mavjud). Mavjud fanlar uchun savollar qo\'shildi (agar kerak bo\'lsa).'))

        self.stdout.write(self.style.SUCCESS('Feyk ma\'lumotlarni yaratish tugallandi!'))